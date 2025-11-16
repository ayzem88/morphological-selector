import re
import ast
import os
import json
import sqlite3
import concurrent.futures
from collections import defaultdict
import multiprocessing
import logging
from datetime import datetime
import hashlib
from pathlib import Path
import pickle
from tqdm import tqdm
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # للاستخدام بدون واجهة رسومية
from jinja2 import Template

# دعم ملفات docx
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("مكتبة python-docx غير مثبتة. ملفات .docx لن تُقرأ. قم بتثبيتها: pip install python-docx")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

##################################
# نظام الكاش
##################################

class CacheManager:
    """نظام كاش لحفظ النتائج المعالجة سابقاً"""
    def __init__(self, cache_dir="cache"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_file = self.cache_dir / "word_cache.json"
        self.cache = self.load_cache()
        
    def load_cache(self):
        """تحميل الكاش من الملف"""
        if self.cache_file.exists():
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def save_cache(self):
        """حفظ الكاش إلى الملف"""
        with open(self.cache_file, 'w', encoding='utf-8') as f:
            json.dump(self.cache, f, ensure_ascii=False, indent=2)
    
    def get_cache_key(self, word, pattern):
        """إنشاء مفتاح فريد للكلمة والوزن"""
        return f"{word}_{pattern}"
    
    def get(self, word, pattern):
        """الحصول على نتيجة من الكاش"""
        key = self.get_cache_key(word, pattern)
        return self.cache.get(key)
    
    def set(self, word, pattern, result):
        """حفظ نتيجة في الكاش"""
        key = self.get_cache_key(word, pattern)
        self.cache[key] = result
        
    def clear(self):
        """مسح الكاش"""
        self.cache = {}
        if self.cache_file.exists():
            self.cache_file.unlink()

##################################
# نظام قاعدة البيانات
##################################

class DatabaseManager:
    """مدير قاعدة البيانات SQLite"""
    def __init__(self, db_path="morphology.db"):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path)
        self.create_tables()
        
    def create_tables(self):
        """إنشاء الجداول"""
        cursor = self.conn.cursor()
        
        # جدول الأوزان
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                pattern TEXT UNIQUE NOT NULL,
                pattern_type TEXT,
                frequency INTEGER DEFAULT 0,
                extra_chars_count INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # جدول الجذور
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS roots (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                root TEXT UNIQUE NOT NULL,
                frequency INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # جدول النتائج
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                word TEXT NOT NULL,
                root_id INTEGER,
                pattern_id INTEGER,
                prefix TEXT,
                suffix TEXT,
                intermediate TEXT,
                frequency INTEGER DEFAULT 1,
                score REAL DEFAULT 0,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (root_id) REFERENCES roots(id),
                FOREIGN KEY (pattern_id) REFERENCES patterns(id),
                UNIQUE(word, root_id, pattern_id)
            )
        ''')
        
        # جدول الإحصائيات
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS statistics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                run_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                total_words INTEGER,
                unique_words INTEGER,
                total_patterns INTEGER,
                total_roots INTEGER,
                success_rate REAL,
                processing_time REAL
            )
        ''')
        
        self.conn.commit()
    
    def insert_pattern(self, pattern, pattern_type=None, extra_chars_count=0):
        """إدراج وزن جديد"""
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT OR IGNORE INTO patterns (pattern, pattern_type, extra_chars_count)
            VALUES (?, ?, ?)
        ''', (pattern, pattern_type, extra_chars_count))
        self.conn.commit()
        return cursor.lastrowid or self.get_pattern_id(pattern)
    
    def insert_root(self, root):
        """إدراج جذر جديد"""
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT OR IGNORE INTO roots (root)
            VALUES (?)
        ''', (root,))
        self.conn.commit()
        return cursor.lastrowid or self.get_root_id(root)
    
    def insert_result(self, word, root, pattern, prefix, suffix, intermediate, score=0):
        """إدراج نتيجة تحليل"""
        root_id = self.insert_root(root)
        pattern_id = self.get_pattern_id(pattern)
        
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO results 
            (word, root_id, pattern_id, prefix, suffix, intermediate, score, frequency)
            VALUES (?, ?, ?, ?, ?, ?, ?,
                COALESCE((SELECT frequency + 1 FROM results 
                         WHERE word = ? AND root_id = ? AND pattern_id = ?), 1))
        ''', (word, root_id, pattern_id, prefix, suffix, intermediate, score,
              word, root_id, pattern_id))
        
        # تحديث تكرار الوزن والجذر
        cursor.execute('UPDATE patterns SET frequency = frequency + 1 WHERE id = ?', (pattern_id,))
        cursor.execute('UPDATE roots SET frequency = frequency + 1 WHERE id = ?', (root_id,))
        
        self.conn.commit()
    
    def get_pattern_id(self, pattern):
        """الحصول على معرف الوزن"""
        cursor = self.conn.cursor()
        cursor.execute('SELECT id FROM patterns WHERE pattern = ?', (pattern,))
        result = cursor.fetchone()
        return result[0] if result else None
    
    def get_root_id(self, root):
        """الحصول على معرف الجذر"""
        cursor = self.conn.cursor()
        cursor.execute('SELECT id FROM roots WHERE root = ?', (root,))
        result = cursor.fetchone()
        return result[0] if result else None
    
    def get_statistics(self):
        """الحصول على الإحصائيات"""
        cursor = self.conn.cursor()
        stats = {}
        
        cursor.execute('SELECT COUNT(*) FROM results')
        stats['total_results'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(DISTINCT word) FROM results')
        stats['unique_words'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM patterns')
        stats['total_patterns'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM roots')
        stats['total_roots'] = cursor.fetchone()[0]
        
        cursor.execute('''
            SELECT pattern, frequency 
            FROM patterns 
            ORDER BY frequency DESC 
            LIMIT 10
        ''')
        stats['top_patterns'] = cursor.fetchall()
        
        cursor.execute('''
            SELECT root, frequency 
            FROM roots 
            ORDER BY frequency DESC 
            LIMIT 10
        ''')
        stats['top_roots'] = cursor.fetchall()
        
        return stats
    
    def close(self):
        """إغلاق الاتصال بقاعدة البيانات"""
        self.conn.close()

##################################
# نظام ترشيح وتقييم الأوزان
##################################

class PatternRanker:
    """نظام ترشيح وتقييم الأوزان المتعددة"""
    def __init__(self, db_manager=None):
        self.extra_chars = set("سأؤئءآإتمونيهىّا")
        self.db_manager = db_manager
        self.pattern_scores = defaultdict(float)
        
    def calculate_score(self, pattern, word, prefix, suffix, results_count):
        """حساب نقاط الوزن"""
        score = 0
        
        # 1. نقاط أحرف الزيادة (الأولوية الأعلى)
        extra_count = sum(1 for c in pattern if c in self.extra_chars)
        score += extra_count * 20
        
        # 2. نقاط التكرار في قاعدة البيانات
        if self.db_manager:
            pattern_id = self.db_manager.get_pattern_id(pattern)
            if pattern_id:
                cursor = self.db_manager.conn.cursor()
                cursor.execute('SELECT frequency FROM patterns WHERE id = ?', (pattern_id,))
                result = cursor.fetchone()
                if result:
                    score += min(result[0] * 0.5, 50)  # حد أقصى 50 نقطة
        
        # 3. نقاط التطابق مع السوابق واللواحق
        if prefix:
            score += 5
        if suffix:
            score += 5
            
        # 4. نقاط نسبة طول الوزن للكلمة
        length_ratio = len(pattern) / len(word) if len(word) > 0 else 0
        if 0.7 <= length_ratio <= 1.3:
            score += 10
            
        # 5. نقاط عدد النتائج المطابقة
        score += min(results_count * 2, 20)  # حد أقصى 20 نقطة
        
        return score
    
    def rank_patterns(self, patterns_results, word):
        """ترتيب الأوزان حسب النقاط"""
        ranked = []
        
        for pattern, results in patterns_results.items():
            total_score = 0
            for prefix, root, suffix in results:
                score = self.calculate_score(pattern, word, prefix, suffix, len(results))
                total_score += score
            
            avg_score = total_score / len(results) if results else 0
            ranked.append((pattern, results, avg_score))
        
        # ترتيب تنازلي حسب النقاط
        ranked.sort(key=lambda x: x[2], reverse=True)
        
        return ranked

##################################
# معالج الدُفعات المتقدم
##################################

class BatchProcessor:
    """معالج دُفعات متقدم للملفات الكبيرة"""
    def __init__(self, chunk_size=1000, save_interval=5000):
        self.chunk_size = chunk_size
        self.save_interval = save_interval
        self.processed_count = 0
        self.checkpoint_file = "processing_checkpoint.pkl"
        
    def save_checkpoint(self, data):
        """حفظ نقطة استعادة"""
        with open(self.checkpoint_file, 'wb') as f:
            pickle.dump({
                'processed_count': self.processed_count,
                'timestamp': datetime.now(),
                'data': data
            }, f)
    
    def load_checkpoint(self):
        """تحميل نقطة الاستعادة"""
        if os.path.exists(self.checkpoint_file):
            with open(self.checkpoint_file, 'rb') as f:
                return pickle.load(f)
        return None
    
    def process_file_in_chunks(self, file_path, process_func):
        """معالجة الملف على دفعات"""
        checkpoint = self.load_checkpoint()
        start_line = 0
        
        if checkpoint:
            response = input(f"تم العثور على نقطة استعادة ({checkpoint['timestamp']}). هل تريد المتابعة من حيث توقفت؟ (y/n): ")
            if response.lower() == 'y':
                start_line = checkpoint['processed_count']
                logging.info(f"متابعة المعالجة من السطر {start_line}")
        
        results = []
        chunk = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            # تخطي الأسطر المعالجة سابقاً
            for _ in range(start_line):
                next(file, None)
            
            # معالجة باقي الملف
            with tqdm(total=sum(1 for _ in open(file_path, 'r', encoding='utf-8')) - start_line,
                     desc="معالجة الملف", unit="سطر") as pbar:
                
                for line_num, line in enumerate(file, start=start_line):
                    chunk.append(line.strip())
                    self.processed_count = line_num
                    
                    # معالجة الدفعة عند الوصول للحجم المحدد
                    if len(chunk) >= self.chunk_size:
                        chunk_results = process_func(chunk)
                        results.extend(chunk_results)
                        chunk = []
                        pbar.update(self.chunk_size)
                    
                    # حفظ نقطة استعادة
                    if line_num % self.save_interval == 0:
                        self.save_checkpoint(results)
                        logging.info(f"تم حفظ نقطة استعادة عند السطر {line_num}")
                
                # معالجة آخر دفعة
                if chunk:
                    chunk_results = process_func(chunk)
                    results.extend(chunk_results)
                    pbar.update(len(chunk))
        
        # حذف ملف نقطة الاستعادة بعد الانتهاء
        if os.path.exists(self.checkpoint_file):
            os.remove(self.checkpoint_file)
        
        return results

##################################
# مولد التقارير
##################################

class ReportGenerator:
    """مولد تقارير HTML وExcel"""
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.report_dir = Path("reports")
        self.report_dir.mkdir(exist_ok=True)
        
    def generate_text_report(self, stats, coverage=None, output_file="report.txt"):
        """توليد تقرير نصي فقط"""
        lines = []
        lines.append("تقرير التحليل الصرفي")
        lines.append("======================")
        lines.append(f"إجمالي النتائج: {stats.get('total_results', 0):,}")
        lines.append(f"الكلمات الفريدة: {stats.get('unique_words', 0):,}")
        lines.append(f"عدد الأوزان: {stats.get('total_patterns', 0):,}")
        lines.append(f"عدد الجذور: {stats.get('total_roots', 0):,}")
        if 'processing_time' in stats:
            lines.append(f"وقت المعالجة: {stats['processing_time']:.2f} ثانية")
        lines.append("")
        if coverage:
            total_words = coverage.get('total_words', 0)
            recognized = coverage.get('recognized', 0)
            unrecognized = coverage.get('unrecognized', 0)
            percent = (recognized / total_words * 100) if total_words else 0.0
            lines.append("ملخص التغطية (الكلمات المقروءة في المدونة)")
            lines.append("-------------------------------------------")
            lines.append(f"إجمالي الكلمات في المدونة: {total_words:,}")
            lines.append(f"المتعرّف عليها: {recognized:,} ({percent:.2f}%)")
            lines.append(f"غير المتعرّف عليها: {unrecognized:,} ({100.0 - percent:.2f}%)")
            if 'recognized_file' in coverage and 'unrecognized_file' in coverage:
                lines.append("")
                lines.append(f"قائمة الكلمات المتعرّف عليها: {coverage['recognized_file']}")
                lines.append(f"قائمة الكلمات غير المتعرّف عليها: {coverage['unrecognized_file']}")
            if 'coverage_html' in coverage:
                lines.append(f"المخطط التفاعلي (HTML): {coverage['coverage_html']}")

        output_path = self.report_dir / output_file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines) + "\n")
        logging.info(f"تم إنشاء التقرير النصي: {output_path}")
        return output_path

    def generate_coverage_outputs(self, all_words_set, recognized_set,
                                   recognized_file_name="recognized_words.txt",
                                   unrecognized_file_name="unrecognized_words.txt",
                                   coverage_html_name="coverage.html"):
        """توليد ملفات التغطية: قوائم الكلمات وملف HTML تفاعلي صغير"""
        # حساب المجموعات
        all_words_sorted = sorted(all_words_set)
        recognized_sorted = sorted(recognized_set & all_words_set)
        unrecognized_sorted = sorted(all_words_set - recognized_set)

        # كتابة الملفات النصية
        recognized_path = self.report_dir / recognized_file_name
        unrecognized_path = self.report_dir / unrecognized_file_name

        with open(recognized_path, 'w', encoding='utf-8') as f:
            for w in recognized_sorted:
                f.write(w + "\n")

        with open(unrecognized_path, 'w', encoding='utf-8') as f:
            for w in unrecognized_sorted:
                f.write(w + "\n")

        # إنشاء HTML تفاعلي بسيط (Canvas) بدون مكتبات خارجية
        recognized_count = len(recognized_sorted)
        unrecognized_count = len(unrecognized_sorted)
        total = max(1, recognized_count + unrecognized_count)

        html_content = f"""
<!DOCTYPE html>
<html dir=\"rtl\" lang=\"ar\">
<head>
  <meta charset=\"UTF-8\" />
  <title>تغطية التعرف على الكلمات</title>
  <style>
    body {{ font-family: Arial, Tahoma, sans-serif; background:#f7f7f7; margin:20px; }}
    .card {{ max-width: 560px; margin: 0 auto; background:#fff; padding:20px; border-radius:12px; box-shadow:0 4px 18px rgba(0,0,0,.08); }}
    h1 {{ margin-top:0; font-size:20px; color:#333; }}
    .legend {{ display:flex; gap:14px; margin:10px 0 0; align-items:center; flex-wrap:wrap; }}
    .legend-item {{ display:flex; gap:8px; align-items:center; font-size:14px; color:#444; }}
    .box {{ width:14px; height:14px; border-radius:3px; }}
    .muted {{ color:#666; font-size:13px; margin-top:6px; }}
    canvas {{ display:block; margin: 10px auto; }}
    .tooltip {{ position:absolute; background:rgba(0,0,0,.8); color:#fff; padding:6px 10px; border-radius:6px; font-size:12px; pointer-events:none; transform:translate(-50%, -140%); white-space:nowrap; }}
  </style>
</head>
<body>
  <div class=\"card\">
    <h1>مخطط التغطية (الكلمات المقروءة)</h1>
    <canvas id=\"chart\" width=\"420\" height=\"420\" aria-label=\"نسبة التعرف\" role=\"img\"></canvas>
    <div class=\"legend\">
      <div class=\"legend-item\"><span class=\"box\" style=\"background:#4CAF50\"></span> متعرّف عليها: {recognized_count} / {total}</div>
      <div class=\"legend-item\"><span class=\"box\" style=\"background:#E53935\"></span> غير متعرّف عليها: {unrecognized_count} / {total}</div>
    </div>
    <div class=\"muted\">حرّك المؤشر فوق المخطط لعرض النسب.</div>
  </div>
  <div id=\"tt\" class=\"tooltip\" style=\"display:none\"></div>
  <script>
    (function(){{
      const recognized = {recognized_count};
      const unknown = {unrecognized_count};
      const total = Math.max(1, recognized + unknown);
      const data = [recognized, unknown];
      const colors = ['#4CAF50', '#E53935'];
      const labels = ['متعرّف عليها', 'غير متعرّف عليها'];

      const canvas = document.getElementById('chart');
      const ctx = canvas.getContext('2d');
      const cx = canvas.width/2, cy = canvas.height/2, r = 150, ir = 90;

      function drawPie(){{
        let start = -Math.PI/2;
        ctx.clearRect(0,0,canvas.width,canvas.height);
        // الخارجي
        for (let i=0;i<data.length;i++){{
          const angle = (data[i]/total) * Math.PI*2;
          ctx.beginPath();
          ctx.moveTo(cx, cy);
          ctx.arc(cx, cy, r, start, start+angle);
          ctx.closePath();
          ctx.fillStyle = colors[i];
          ctx.fill();
          start += angle;
        }}
        // ثقب داخلي (دونات)
        ctx.globalCompositeOperation = 'destination-out';
        ctx.beginPath();
        ctx.arc(cx, cy, ir, 0, Math.PI*2);
        ctx.fill();
        ctx.globalCompositeOperation = 'source-over';

        // نص الوسط
        const percent = Math.round((recognized/total)*1000)/10;
        ctx.fillStyle = '#333';
        ctx.font = 'bold 22px Tahoma, Arial';
        ctx.textAlign = 'center';
        ctx.textBaseline = 'middle';
        ctx.fillText(percent + '%', cx, cy-6);
        ctx.font = '13px Tahoma, Arial';
        ctx.fillText('متعرّف عليها', cx, cy+14);
      }}

      function hitTest(x, y){{
        const dx = x - cx, dy = y - cy; const d = Math.sqrt(dx*dx + dy*dy);
        if (d < ir || d > r) return -1;
        let angle = Math.atan2(dy, dx);
        if (angle < -Math.PI/2) angle += Math.PI*2; // محاذاة البدء
        let acc = -Math.PI/2;
        for (let i=0;i<data.length;i++){{
          const a = (data[i]/total)*Math.PI*2;
          if (angle >= acc && angle < acc + a) return i;
          acc += a;
        }}
        return -1;
      }}

      const tooltip = document.getElementById('tt');
      canvas.addEventListener('mousemove', (e)=>{{
        const rect = canvas.getBoundingClientRect();
        const x = e.clientX - rect.left; const y = e.clientY - rect.top;
        const i = hitTest(x, y);
        if (i === -1) {{ tooltip.style.display = 'none'; return; }}
        const value = data[i];
        const pct = Math.round((value/total)*1000)/10;
        tooltip.style.display = 'block';
        tooltip.textContent = labels[i] + ': ' + value + ' (' + pct + '%)';
        tooltip.style.left = (e.pageX) + 'px';
        tooltip.style.top = (e.pageY) + 'px';
      }});
      canvas.addEventListener('mouseleave', ()=>{{ tooltip.style.display='none'; }});

      drawPie();
    }})();
  </script>
  </body>
  </html>
        """

        coverage_path = self.report_dir / coverage_html_name
        with open(coverage_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        logging.info(f"تم إنشاء ملفات التغطية: {recognized_path}, {unrecognized_path}, {coverage_path}")

        return {
            'total_words': len(all_words_sorted),
            'recognized': recognized_count,
            'unrecognized': unrecognized_count,
            'recognized_file': str(recognized_path),
            'unrecognized_file': str(unrecognized_path),
            'coverage_html': str(coverage_path)
        }

    def generate_html_report(self, stats, output_file="report.html"):
        """توليد تقرير HTML"""
        template_str = '''
        <!DOCTYPE html>
        <html dir="rtl" lang="ar">
        <head>
            <meta charset="UTF-8">
            <title>تقرير التحليل الصرفي</title>
            <style>
                body {
                    font-family: 'Arial', 'Tahoma', sans-serif;
                    margin: 20px;
                    background-color: #f5f5f5;
                }
                .container {
                    max-width: 1200px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 20px;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }
                h1, h2 {
                    color: #333;
                    border-bottom: 2px solid #4CAF50;
                    padding-bottom: 10px;
                }
                .stats-grid {
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                    gap: 20px;
                    margin: 20px 0;
                }
                .stat-card {
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 20px;
                    border-radius: 10px;
                    text-align: center;
                }
                .stat-number {
                    font-size: 2em;
                    font-weight: bold;
                    margin: 10px 0;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }
                th, td {
                    padding: 12px;
                    text-align: right;
                    border-bottom: 1px solid #ddd;
                }
                th {
                    background-color: #4CAF50;
                    color: white;
                }
                tr:hover {
                    background-color: #f5f5f5;
                }
                .chart-container {
                    margin: 20px 0;
                    text-align: center;
                }
                .timestamp {
                    text-align: center;
                    color: #666;
                    margin-top: 20px;
                }
                .progress-bar {
                    background-color: #f0f0f0;
                    border-radius: 10px;
                    overflow: hidden;
                    margin: 10px 0;
                }
                .progress-fill {
                    background: linear-gradient(90deg, #4CAF50, #45a049);
                    height: 30px;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    color: white;
                    font-weight: bold;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>📊 تقرير التحليل الصرفي</h1>
                
                <div class="stats-grid">
                    <div class="stat-card">
                        <div>إجمالي النتائج</div>
                        <div class="stat-number">{{ total_results }}</div>
                    </div>
                    <div class="stat-card">
                        <div>الكلمات الفريدة</div>
                        <div class="stat-number">{{ unique_words }}</div>
                    </div>
                    <div class="stat-card">
                        <div>عدد الأوزان</div>
                        <div class="stat-number">{{ total_patterns }}</div>
                    </div>
                    <div class="stat-card">
                        <div>عدد الجذور</div>
                        <div class="stat-number">{{ total_roots }}</div>
                    </div>
                </div>
                
                <h2>🏆 الأوزان الأكثر شيوعاً</h2>
                <table>
                    <thead>
                        <tr>
                            <th>الترتيب</th>
                            <th>الوزن</th>
                            <th>التكرار</th>
                            <th>النسبة المئوية</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for i, (pattern, freq) in enumerate(top_patterns, 1) %}
                        <tr>
                            <td>{{ i }}</td>
                            <td style="font-weight: bold;">{{ pattern }}</td>
                            <td>{{ freq }}</td>
                            <td>
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: {{ (freq/max_pattern_freq)*100 }}%;">
                                        {{ "%.1f"|format((freq/total_pattern_freq)*100) }}%
                                    </div>
                                </div>
                            </td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
                
                <h2>🌳 الجذور الأكثر شيوعاً</h2>
                <table>
                    <thead>
                        <tr>
                            <th>الترتيب</th>
                            <th>الجذر</th>
                            <th>التكرار</th>
                            <th>النسبة المئوية</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for i, (root, freq) in enumerate(top_roots, 1) %}
                        <tr>
                            <td>{{ i }}</td>
                            <td style="font-weight: bold;">{{ root }}</td>
                            <td>{{ freq }}</td>
                            <td>
                                <div class="progress-bar">
                                    <div class="progress-fill" style="width: {{ (freq/max_root_freq)*100 }}%;">
                                        {{ "%.1f"|format((freq/total_root_freq)*100) }}%
                                    </div>
                                </div>
                            </td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
                
                <div class="chart-container">
                    <h2>📈 الرسوم البيانية</h2>
                    <img src="patterns_chart.png" alt="توزيع الأوزان" style="max-width: 100%;">
                    <img src="roots_chart.png" alt="توزيع الجذور" style="max-width: 100%;">
                </div>
                
                <div class="timestamp">
                    تم إنشاء التقرير: {{ timestamp }}
                </div>
            </div>
        </body>
        </html>
        '''
        
        template = Template(template_str)
        
        # حساب الإحصائيات الإضافية
        max_pattern_freq = max([f for _, f in stats['top_patterns']]) if stats['top_patterns'] else 1
        total_pattern_freq = sum([f for _, f in stats['top_patterns']])
        max_root_freq = max([f for _, f in stats['top_roots']]) if stats['top_roots'] else 1
        total_root_freq = sum([f for _, f in stats['top_roots']])
        
        html_content = template.render(
            total_results=stats['total_results'],
            unique_words=stats['unique_words'],
            total_patterns=stats['total_patterns'],
            total_roots=stats['total_roots'],
            top_patterns=stats['top_patterns'],
            top_roots=stats['top_roots'],
            max_pattern_freq=max_pattern_freq,
            total_pattern_freq=total_pattern_freq,
            max_root_freq=max_root_freq,
            total_root_freq=total_root_freq,
            timestamp=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            enumerate=enumerate
        )
        
        output_path = self.report_dir / output_file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # توليد الرسوم البيانية
        self.generate_charts(stats)
        
        logging.info(f"تم إنشاء التقرير HTML: {output_path}")
        return output_path
    
    def generate_charts(self, stats):
        """توليد الرسوم البيانية"""
        plt.rcParams['font.family'] = ['Arial Unicode MS', 'Tahoma']
        
        # رسم بياني للأوزان
        if stats['top_patterns']:
            patterns, frequencies = zip(*stats['top_patterns'])
            
            plt.figure(figsize=(12, 6))
            plt.bar(range(len(patterns)), frequencies, color='#4CAF50')
            plt.xlabel('الوزن', fontsize=12)
            plt.ylabel('التكرار', fontsize=12)
            plt.title('الأوزان الأكثر شيوعاً', fontsize=14, fontweight='bold')
            plt.xticks(range(len(patterns)), patterns, rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(self.report_dir / 'patterns_chart.png', dpi=100, bbox_inches='tight')
            plt.close()
        
        # رسم بياني للجذور
        if stats['top_roots']:
            roots, frequencies = zip(*stats['top_roots'])
            
            plt.figure(figsize=(12, 6))
            plt.bar(range(len(roots)), frequencies, color='#2196F3')
            plt.xlabel('الجذر', fontsize=12)
            plt.ylabel('التكرار', fontsize=12)
            plt.title('الجذور الأكثر شيوعاً', fontsize=14, fontweight='bold')
            plt.xticks(range(len(roots)), roots, rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(self.report_dir / 'roots_chart.png', dpi=100, bbox_inches='tight')
            plt.close()
    
    def generate_excel_report(self, stats, output_file="report.xlsx"):
        """توليد تقرير Excel"""
        output_path = self.report_dir / output_file
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # ورقة الإحصائيات العامة
            summary_data = {
                'المؤشر': ['إجمالي النتائج', 'الكلمات الفريدة', 'عدد الأوزان', 'عدد الجذور'],
                'القيمة': [stats['total_results'], stats['unique_words'], 
                          stats['total_patterns'], stats['total_roots']]
            }
            df_summary = pd.DataFrame(summary_data)
            df_summary.to_excel(writer, sheet_name='ملخص', index=False)
            
            # ورقة الأوزان
            if stats['top_patterns']:
                patterns_data = {
                    'الوزن': [p for p, _ in stats['top_patterns']],
                    'التكرار': [f for _, f in stats['top_patterns']]
                }
                df_patterns = pd.DataFrame(patterns_data)
                df_patterns.to_excel(writer, sheet_name='الأوزان', index=False)
            
            # ورقة الجذور
            if stats['top_roots']:
                roots_data = {
                    'الجذر': [r for r, _ in stats['top_roots']],
                    'التكرار': [f for _, f in stats['top_roots']]
                }
                df_roots = pd.DataFrame(roots_data)
                df_roots.to_excel(writer, sheet_name='الجذور', index=False)
        
        logging.info(f"تم إنشاء تقرير Excel: {output_path}")
        return output_path

##################################
# نظام التحقق التبادلي
##################################

class CrossValidator:
    """نظام التحقق من صحة التحليل"""
    def __init__(self):
        self.validation_results = []
        
    def reconstruct_word(self, root, pattern, prefix="", suffix=""):
        """إعادة بناء الكلمة من المكونات"""
        # استبدال أحرف الوزن (فعل) بأحرف الجذر
        reconstructed = pattern
        root_chars = list(root)
        pattern_chars = ['ف', 'ع', 'ل']
        
        for i, char in enumerate(pattern_chars[:len(root_chars)]):
            if char in reconstructed:
                reconstructed = reconstructed.replace(char, root_chars[i], 1)
        
        # إضافة السوابق واللواحق
        return prefix + reconstructed + suffix
    
    def validate_analysis(self, original_word, root, pattern, prefix, suffix):
        """التحقق من صحة التحليل"""
        reconstructed = self.reconstruct_word(root, pattern, prefix, suffix)
        
        # حساب نسبة التطابق
        match_ratio = self.calculate_similarity(original_word, reconstructed)
        
        validation_result = {
            'original': original_word,
            'reconstructed': reconstructed,
            'root': root,
            'pattern': pattern,
            'prefix': prefix,
            'suffix': suffix,
            'match_ratio': match_ratio,
            'is_valid': match_ratio > 0.8  # عتبة 80% للصحة
        }
        
        self.validation_results.append(validation_result)
        return validation_result
    
    def calculate_similarity(self, word1, word2):
        """حساب نسبة التشابه بين كلمتين"""
        # إزالة التشكيل للمقارنة
        word1_clean = re.sub(r'[ًٌٍَُِّْ]', '', word1)
        word2_clean = re.sub(r'[ًٌٍَُِّْ]', '', word2)
        
        if word1_clean == word2_clean:
            return 1.0
        
        # حساب نسبة التشابه باستخدام Levenshtein distance
        max_len = max(len(word1_clean), len(word2_clean))
        if max_len == 0:
            return 0.0
        
        distance = self.levenshtein_distance(word1_clean, word2_clean)
        return 1.0 - (distance / max_len)
    
    def levenshtein_distance(self, s1, s2):
        """حساب مسافة Levenshtein"""
        if len(s1) < len(s2):
            return self.levenshtein_distance(s2, s1)
        
        if len(s2) == 0:
            return len(s1)
        
        previous_row = range(len(s2) + 1)
        for i, c1 in enumerate(s1):
            current_row = [i + 1]
            for j, c2 in enumerate(s2):
                insertions = previous_row[j + 1] + 1
                deletions = current_row[j] + 1  # كان ناقص الرقم 1
                substitutions = previous_row[j] + (c1 != c2)
                current_row.append(min(insertions, deletions, substitutions))
            previous_row = current_row
        
        return previous_row[-1]
    
    def get_validation_report(self):
        """الحصول على تقرير التحقق"""
        if not self.validation_results:
            return "لا توجد نتائج للتحقق"
        
        total = len(self.validation_results)
        valid = sum(1 for r in self.validation_results if r['is_valid'])
        invalid = total - valid
        
        report = f"""
        تقرير التحقق التبادلي:
        =====================
        إجمالي الكلمات المحللة: {total}
        التحليلات الصحيحة: {valid} ({valid/total*100:.1f}%)
        التحليلات غير الصحيحة: {invalid} ({invalid/total*100:.1f}%)
        
        أمثلة على التحليلات غير الصحيحة:
        """
        
        for result in self.validation_results[:10]:  # أول 10 أخطاء
            if not result['is_valid']:
                report += f"""
        الكلمة الأصلية: {result['original']}
        الكلمة المعاد بناؤها: {result['reconstructed']}
        نسبة التطابق: {result['match_ratio']:.1%}
        """
        
        return report
##################################
# محسّن الذاكرة
##################################

class MemoryOptimizer:
   """محسّن استخدام الذاكرة للملفات الضخمة"""
   def __init__(self, max_memory_mb=500):
       self.max_memory_mb = max_memory_mb
       self.current_memory = 0
       
   def process_file_stream(self, file_path, process_func, batch_size=100):
       """معالجة الملف بطريقة التدفق"""
       def file_generator():
           with open(file_path, 'r', encoding='utf-8') as f:
               batch = []
               for line in f:
                   batch.append(line.strip())
                   if len(batch) >= batch_size:
                       yield batch
                       batch = []
               if batch:
                   yield batch
       
       # معالجة الدفعات واحدة تلو الأخرى
       for batch in file_generator():
           results = process_func(batch)
           # معالجة النتائج وحذفها من الذاكرة
           yield results
           del results
           del batch

##################################
# فئات المعالجة اللغوية المحسّنة
##################################

class ArabicProcessor:
   def __init__(self, optional_tashkeel=False, symbols_map=None, cache_manager=None):
       self.arabic_diacritics_pattern = re.compile("[ًٌٍَُِّْ]")
       self.optional_tashkeel = optional_tashkeel
       self.arabic_symbols = symbols_map if symbols_map else {}
       self.cache_manager = cache_manager

   def add_optional_tashkeel_and_grouping(self, pattern):
       if self.optional_tashkeel:
           return re.sub(self.arabic_diacritics_pattern, lambda m: f"[{m.group()}]?", pattern)
       else:
           return re.sub(self.arabic_diacritics_pattern, lambda m: m.group(), pattern)

   def replace_symbols(self, word):
       return ''.join(self.arabic_symbols.get(letter, letter) for letter in word)


class DiacriticsHandler:
   DIACRITICS = 'ًٌٍَُِّْْٰ'

   @staticmethod
   def remove_diacritics(word):
       return ''.join(c for c in word if c not in DiacriticsHandler.DIACRITICS)

   @staticmethod
   def group_letters_with_diacritics(word):
       letters_with_diacritics = []
       i = 0
       while i < len(word):
           c = word[i]
           if c not in DiacriticsHandler.DIACRITICS:
               letter = c
               i += 1
               while i < len(word) and word[i] in DiacriticsHandler.DIACRITICS:
                   letter += word[i]
                   i += 1
               letters_with_diacritics.append(letter)
           else:
               i += 1
       return letters_with_diacritics

   @staticmethod
   def normalize_quranic_text(word):
       """تطبيع النص القرآني - توحيد الحروف والهمزات"""
       # توحيد الهمزات
       word = re.sub(r'[ءأإآ]', 'أ', word)
       
       # توحيد الألفات
       word = re.sub(r'[اٱ]', 'ا', word)
       
       # إزالة الحروف القرآنية الخاصة
       word = re.sub(r'[ٰٱٲٳٴٵٶٷٸٹٺٻټٽپٿۖۗۘۙۚۛۜ۝۞ۣ۟۠ۡۢۤۥۦۧۨ۩۪ۭ۫۬ۮۯ]', '', word)
       
       # توحيد التاءات
       word = re.sub(r'ة', 'ت', word)
       
       return word


class WordSplitter:
   ROOT_INDICATORS = 'فعل'

   def __init__(self, diacritics_handler):
       self.diacritics_handler = diacritics_handler

   def split_word(self, template_word, target_word):
       template_clean = self.diacritics_handler.remove_diacritics(template_word)
       template_letters = list(template_clean)
       target_letters = self.diacritics_handler.group_letters_with_diacritics(target_word)

       if len(template_letters) != len(target_letters):
           return '', '', target_word, ''

       root_positions = [idx for idx, c in enumerate(template_letters) if c in WordSplitter.ROOT_INDICATORS]
       if not root_positions:
           return '', '', target_word, ''

       prefix_letters = []
       root_letters = []
       intermediate_letters = []
       suffix_letters = []

       first_root_pos = root_positions[0]
       last_root_pos = root_positions[-1]

       for idx, c in enumerate(template_letters):
           target_char = target_letters[idx]
           if c in WordSplitter.ROOT_INDICATORS:
               root_letters.append(target_char)
           else:
               if idx < first_root_pos:
                   prefix_letters.append(target_char)
               elif idx > last_root_pos:
                   suffix_letters.append(target_char)
               else:
                   intermediate_letters.append(target_char)

       prefix = ''.join(prefix_letters)
       intermediate = ''.join(intermediate_letters)
       root = ''.join(root_letters)
       suffix = ''.join(suffix_letters)

       return prefix, intermediate, root, suffix

##################################
# دالة اكتشاف نوع الملف تلقائياً
##################################

def detect_file_type(file_path, sample_lines=50):
    """
    اكتشاف نوع الملف تلقائياً: 'list' أو 'text'
    
    المعايير المحسّنة:
    - إذا كان متوسط عدد الكلمات في السطر > 2.0 → 'text'
    - إذا كان متوسط عدد الكلمات في السطر <= 2.0 → 'list'
    - إذا كان أكثر من 80% من الأسطر تحتوي على كلمة واحدة فقط → 'list'
    - إذا كان أكثر من 50% من الأسطر تحتوي على أكثر من 3 كلمات → 'text'
    """
    arabic_word_re = re.compile(r"[\u0621-\u064A\u064B-\u065F\u0670\u06D6-\u06DC\u06DF-\u06E8\u06EA-\u06ED]+")
    word_counts = []
    single_word_lines = 0
    multi_word_lines = 0
    
    file_ext = os.path.splitext(file_path)[1].lower()
    lines = []
    
    try:
        if file_ext == '.docx':
            # قراءة ملف docx
            if not DOCX_SUPPORT:
                logging.warning(f"مكتبة python-docx غير مثبتة. استخدام 'text' كافتراضي للملف: {file_path}")
                return 'text'
            
            try:
                doc = Document(file_path)
                for i, paragraph in enumerate(doc.paragraphs):
                    if i >= sample_lines:
                        break
                    text = paragraph.text.strip()
                    if text:
                        lines.append(text)
            except Exception as e:
                logging.warning(f"خطأ في قراءة ملف docx {file_path}: {e}. استخدام 'text' كافتراضي.")
                return 'text'
        else:
            # قراءة ملف txt
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    for i, line in enumerate(f):
                        if i >= sample_lines:
                            break
                        line = line.strip()
                        if line:
                            lines.append(line)
            except Exception as e:
                logging.warning(f"خطأ في اكتشاف نوع الملف {file_path}: {e}. استخدام 'text' كافتراضي.")
                return 'text'  # افتراضي في حالة الخطأ
        
        # تحليل الأسطر
        for line in lines:
            words = arabic_word_re.findall(line)
            word_count = len(words)
            if word_count > 0:
                word_counts.append(word_count)
                if word_count == 1:
                    single_word_lines += 1
                elif word_count > 3:
                    multi_word_lines += 1
    except Exception as e:
        logging.warning(f"خطأ في اكتشاف نوع الملف {file_path}: {e}. استخدام 'text' كافتراضي.")
        return 'text'  # افتراضي في حالة الخطأ
    
    if not word_counts:
        logging.warning(f"الملف {file_path} فارغ أو لا يحتوي على كلمات عربية. استخدام 'text' كافتراضي.")
        return 'text'  # افتراضي للملفات الفارغة
    
    total_lines = len(word_counts)
    avg_words = sum(word_counts) / total_lines
    single_word_ratio = single_word_lines / total_lines if total_lines > 0 else 0
    multi_word_ratio = multi_word_lines / total_lines if total_lines > 0 else 0
    
    # منطق محسّن للاكتشاف
    if single_word_ratio > 0.8:
        # أكثر من 80% من الأسطر تحتوي على كلمة واحدة → قائمة
        detected_type = 'list'
    elif multi_word_ratio > 0.5:
        # أكثر من 50% من الأسطر تحتوي على أكثر من 3 كلمات → نص
        detected_type = 'text'
    elif avg_words > 2.0:
        # متوسط أكثر من 2 كلمة → نص
        detected_type = 'text'
    else:
        # متوسط أقل من أو يساوي 2 كلمة → قائمة
        detected_type = 'list'
    
    logging.debug(f"تم اكتشاف نوع الملف {file_path}: {detected_type} (متوسط الكلمات: {avg_words:.2f}, "
                  f"نسبة السطور بكلمة واحدة: {single_word_ratio:.2%}, "
                  f"نسبة السطور بأكثر من 3 كلمات: {multi_word_ratio:.2%})")
    return detected_type

##################################
# فئة إدارة الملفات والبحث المحسّنة
##################################

class FileManager:
   EXTRA_CHARS = "سأؤئءآإتمونيهىّا"

   def __init__(self, corpus_type='text', match_whole_word=True, affixes_data=None, 
                tags_map=None, db_manager=None, cache_manager=None, 
                pattern_ranker=None, cross_validator=None):
       self.corpus_type = corpus_type
       self.match_whole_word = match_whole_word
       affixes_data = affixes_data if affixes_data else {'prefixes': [], 'suffixes': []}

       self.prefixes = affixes_data.get('prefixes', [])
       self.suffixes = affixes_data.get('suffixes', [])

       self.prefix_pattern = f"(?:{'|'.join(map(re.escape, self.prefixes))})" if self.prefixes else ""
       self.suffix_pattern = f"(?:{'|'.join(map(re.escape, self.suffixes))})" if self.suffixes else ""

       # بناء word_boundary حسب corpus_type (كما في الكود الأصلي)
       if self.corpus_type == 'list':
           if self.match_whole_word:
               self.word_boundary_start = r"^"
               self.word_boundary_end = r"$"
           else:
               self.word_boundary_start = ''
               self.word_boundary_end = ''
       elif self.corpus_type == 'text':
           if self.match_whole_word:
               self.word_boundary = r"(?:\s|^|$|[،؛.!؟\"'«»()\-])"
               # إضافة حدود الكلمة الكاملة للنص
               self.word_boundary_start = r"(?:\s|^|[،؛.!؟\"'«»()\-])"
               self.word_boundary_end = r"(?:\s|$|[،؛.!؟\"'«»()\-])"
           else:
               self.word_boundary = ''
               self.word_boundary_start = ''
               self.word_boundary_end = ''
       else:
           self.word_boundary = ''

       self.diacritics_handler = DiacriticsHandler()
       self.word_splitter = WordSplitter(self.diacritics_handler)
       self.tags_map = tags_map if tags_map else {}
       self.db_manager = db_manager
       self.cache_manager = cache_manager
       self.pattern_ranker = pattern_ranker
       self.cross_validator = cross_validator

   def search_patterns_in_file(self, file_path, pattern, weight):
       """البحث عن الأنماط في الملف مع استخدام الكاش والتطبيع"""
       # التحقق من الكاش أولاً
       if self.cache_manager:
           cache_key = f"{file_path}_{pattern}"
           cached_result = self.cache_manager.get(cache_key, pattern)
           if cached_result:
               logging.info(f"تم العثور على نتيجة في الكاش للنمط: {pattern}")
               return cached_result
       
       # بناء الأنماط حسب corpus_type (كما في الكود الأصلي)
       results = []
       if self.corpus_type == 'list':
           full_pattern = f"{self.word_boundary_start}(?P<prefix>{self.prefix_pattern})?(?P<root>{pattern})(?P<suffix>{self.suffix_pattern})?{self.word_boundary_end}"
       elif self.corpus_type == 'text':
           if self.match_whole_word:
               # استخدام حدود الكلمة الكاملة للنص
               full_pattern = f"{self.word_boundary_start}(?P<prefix>{self.prefix_pattern})?(?P<root>{pattern})(?P<suffix>{self.suffix_pattern})?{self.word_boundary_end}"
           else:
               full_pattern = f"{self.word_boundary}(?P<prefix>{self.prefix_pattern})?(?P<root>{pattern})(?P<suffix>{self.suffix_pattern})?{self.word_boundary}"
       else:
           full_pattern = f"(?P<prefix>{self.prefix_pattern})?(?P<root>{pattern})(?P<suffix>{self.suffix_pattern})?"

       compiled_pattern = re.compile(full_pattern)
       logging.debug(f"استخدام النمط: {compiled_pattern.pattern}")

       # قراءة الملف حسب نوعه
       file_ext = os.path.splitext(file_path)[1].lower()
       
       if file_ext == '.docx':
           # قراءة ملف docx
           if not DOCX_SUPPORT:
               logging.warning(f"مكتبة python-docx غير مثبتة. تخطي الملف: {file_path}")
               return results
           
           try:
               doc = Document(file_path)
               lines = []
               for paragraph in doc.paragraphs:
                   if paragraph.text.strip():
                       lines.append(paragraph.text)
           except Exception as e:
               logging.error(f"خطأ في قراءة ملف docx {file_path}: {e}")
               return results
           
           # معالجة محتوى docx
           for line in lines:
               # تطبيع السطر قبل البحث
               normalized_line = self.diacritics_handler.normalize_quranic_text(line)
               
               for match in compiled_pattern.finditer(normalized_line):
                   prefix = match.group('prefix') or ''
                   root = match.group('root')
                   suffix = match.group('suffix') or ''
                   
                   # التحقق التبادلي
                   if self.cross_validator:
                       word = prefix + root + suffix
                       validation = self.cross_validator.validate_analysis(
                           word, root, pattern, prefix, suffix
                       )
                       if validation['is_valid']:
                           results.append((prefix, root, suffix))
                   else:
                       results.append((prefix, root, suffix))
       else:
           # قراءة ملف txt سطراً بسطر (كما في الكود الأصلي)
           with open(file_path, 'r', encoding='utf-8') as file:
               for line in file:
                   # تطبيع السطر قبل البحث
                   normalized_line = self.diacritics_handler.normalize_quranic_text(line)
                   
                   for match in compiled_pattern.finditer(normalized_line):
                       prefix = match.group('prefix') or ''
                       root = match.group('root')
                       suffix = match.group('suffix') or ''
                       
                       # التحقق التبادلي
                       if self.cross_validator:
                           word = prefix + root + suffix
                           validation = self.cross_validator.validate_analysis(
                               word, root, pattern, prefix, suffix
                           )
                           if validation['is_valid']:
                               results.append((prefix, root, suffix))
                       else:
                           results.append((prefix, root, suffix))
       
       # حفظ في الكاش
       if self.cache_manager and results:
           cache_key = f"{file_path}_{pattern}"
           self.cache_manager.set(cache_key, pattern, results)
       
       return results

   def _count_results(self, results):
       counts = defaultdict(int)
       for prefix, root, suffix in results:
           matched_word = prefix + root + suffix
           counts[(matched_word, prefix, root, suffix)] += 1
       return counts

   def write_results(self, folder_path, weight, results):
       """كتابة النتائج مع الحفظ في قاعدة البيانات"""
       if not results:
           logging.info(f"لا توجد نتائج للوزن: {weight}. لن يتم إنشاء ملف.")
           return

       if not os.path.exists(folder_path):
           os.makedirs(folder_path, exist_ok=True)

       # حساب عدد أحرف الزيادة للوزن
       extra_chars_count = sum(1 for c in weight if c in FileManager.EXTRA_CHARS)
       
       # إدراج الوزن في قاعدة البيانات
       if self.db_manager:
           pattern_type = 'اسم' if 'الأسماء' in folder_path else 'فعل'
           self.db_manager.insert_pattern(weight, pattern_type, extra_chars_count)

       file_path = os.path.join(folder_path, f"{weight}.txt")
       with open(file_path, 'w', encoding='utf-8') as file:
           for (matched_word, prefix, root, suffix), count in self._count_results(results).items():
               original_word = matched_word
               template_word = weight
               target_word = root

               prefix_morph, intermediate_morph, root_morph, suffix_morph = self.word_splitter.split_word(
                   template_word, target_word
               )
               root_without_diacritics = self.diacritics_handler.remove_diacritics(root_morph)

               prefix_output = f"[{prefix if prefix else '#'}]"
               suffix_output = f"[{suffix if suffix else '#'}]"
               intermediate_output = f"[{intermediate_morph if intermediate_morph else '#'}]"

               result_line = (f"{root_without_diacritics} | {target_word} | {template_word} | {original_word} | "
                              f"{prefix_output} | {intermediate_output} | {suffix_output} | تكرار: {count}")

               # إضافة معلومات الوسم إن وجدت
               if template_word in self.tags_map:
                   result_line += f" | [الوسم: {self.tags_map[template_word]}]"

               file.write(f"{result_line}\n")
               
               # حفظ في قاعدة البيانات
               if self.db_manager:
                   # حساب النقاط للنتيجة
                   score = 0
                   if self.pattern_ranker:
                       score = self.pattern_ranker.calculate_score(
                           weight, matched_word, prefix, suffix, count
                       )
                   
                   self.db_manager.insert_result(
                       matched_word, root_without_diacritics, weight,
                       prefix, suffix, intermediate_morph, score
                   )

   def read_weights_and_derived_words(self, file_path):
       weights = {}
       with open(file_path, 'r', encoding='utf-8') as file:
           for line in file:
               if not line.startswith('#'):
                   parts = line.strip().split(':')
                   if len(parts) == 2:
                       key, values = parts[0].strip(), parts[1].strip()
                       derived_words = [value.strip() for value in values.split('،')]
                       weights[key] = derived_words
                   else:
                       if parts:
                           weights[parts[0].strip()] = []

       # إعادة ترتيب الأوزان بناءً على عدد أحرف الزيادة
       weights = self._reorder_weights(weights)
       return weights

   def _reorder_weights(self, weights_dict):
       # أحرف الزيادة المقترحة
       extra_chars = set(FileManager.EXTRA_CHARS)

       def count_extra_chars(word):
           # يحسب عدد الأحرف الزائدة في الوزن
           return sum(letter in extra_chars for letter in word)

       # تحويل الدكت إلى قائمة من tuples: (weight, derived_words)
       weights_list = [(w, d) for w, d in weights_dict.items()]

       # ترتيب القائمة حسب عدد أحرف الزيادة تنازلياً
       weights_list.sort(key=lambda x: count_extra_chars(x[0]), reverse=True)

       # إعادة تحويلها إلى dict بنفس الترتيب الجديد
       new_weights_dict = {}
       for w, d in weights_list:
           new_weights_dict[w] = d
       return new_weights_dict

##################################
# وظائف المعالجة الرئيسية المحسّنة
##################################

def process_weight(args):
   weight, derived_weights, file_paths, results_dir_name, corpus_type, match_whole_word, affixes_data, tags_map, symbols_map, optional_tashkeel, use_cross_validation = args
   logging.info(f"بدأ معالجة الوزن: {weight}")
   
   # إنشاء الكائنات المطلوبة داخل كل عملية
   processor = ArabicProcessor(
       optional_tashkeel=optional_tashkeel, 
       symbols_map=symbols_map,
       cache_manager=None  # لا نستخدم الكاش في multiprocessing
   )
   
   # في multiprocessing، لا نستخدم cross_validator (كما في الكود الأصلي)
   # cross_validator يتم استخدامه فقط في الكود الرئيسي بعد جمع النتائج
   file_manager = FileManager(
       corpus_type=corpus_type, 
       match_whole_word=match_whole_word, 
       affixes_data=affixes_data, 
       tags_map=tags_map,
       db_manager=None,  # لا نستخدم قاعدة البيانات في multiprocessing
       cache_manager=None,
       pattern_ranker=None,
       cross_validator=None  # لا نستخدم cross_validator في multiprocessing (كما في الكود الأصلي)
   )
   
   # تطبيع الوزن قبل المعالجة
   diacritics_handler = DiacriticsHandler()
   normalized_weight = diacritics_handler.normalize_quranic_text(weight)
   
   pattern = processor.add_optional_tashkeel_and_grouping(weight)
   pattern = processor.replace_symbols(pattern)

   all_results = []
   patterns_results = defaultdict(list)  # لتجميع النتائج حسب الوزن
   
   for file_path in file_paths:
       logging.info(f"معالجة الملف: {file_path} للوزن: {weight}")
       found_results = file_manager.search_patterns_in_file(file_path, pattern, weight)
       all_results.extend(found_results)
       patterns_results[weight].extend(found_results)

   folder_path = os.path.join(results_dir_name, weight)
   file_manager.write_results(folder_path, weight, all_results)

   # معالجة الأوزان المشتقة
   for derived_weight in derived_weights:
       logging.info(f"بدأ معالجة الوزن المشتق: {derived_weight} للوزن الأساسي: {weight}")
       derived_pattern = processor.add_optional_tashkeel_and_grouping(derived_weight)
       derived_pattern = processor.replace_symbols(derived_pattern)

       all_derived_results = []
       for file_path in file_paths:
           logging.info(f"معالجة الملف: {file_path} للوزن المشتق: {derived_weight}")
           found_results = file_manager.search_patterns_in_file(file_path, derived_pattern, derived_weight)
           all_derived_results.extend(found_results)
           patterns_results[derived_weight].extend(found_results)

       file_manager.write_results(folder_path, derived_weight, all_derived_results)

   logging.info(f"انتهى معالجة الوزن: {weight}")
   # إرجاع البيانات للحفظ في قاعدة البيانات
   return {
       'weight': weight,
       'results': all_results,
       'patterns_results': dict(patterns_results),
       'count': len(all_results),
       'validation_results': []  # لا نستخدم cross_validator في multiprocessing
   }

##################################
# تجميع كلمات المدونة
##################################

def collect_corpus_words(file_paths, corpus_type='list'):
    """تجميع كل الكلمات الواردة في المدونة (مع إزالة التشكيل والتطبيع)."""
    diacritics_handler = DiacriticsHandler()
    all_words = set()
    # نمط استخراج الكلمات العربية (يشمل التشكيل)
    arabic_word_re = re.compile(r"[\u0621-\u064A\u064B-\u065F\u0670\u06D6-\u06DC\u06DF-\u06E8\u06EA-\u06ED]+")

    for fp in file_paths:
        try:
            # اكتشاف نوع الملف تلقائياً
            file_type = detect_file_type(fp)
            
            # قراءة الملف حسب نوعه
            file_ext = os.path.splitext(fp)[1].lower()
            lines = []
            
            if file_ext == '.docx':
                # قراءة ملف docx
                if not DOCX_SUPPORT:
                    logging.warning(f"مكتبة python-docx غير مثبتة. تخطي الملف: {fp}")
                    continue
                
                try:
                    doc = Document(fp)
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            lines.append(paragraph.text)
                except Exception as e:
                    logging.warning(f"خطأ في قراءة ملف docx {fp}: {e}")
                    continue
            else:
                # قراءة ملف txt
                try:
                    with open(fp, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                except Exception as e:
                    logging.warning(f"خطأ في قراءة الملف {fp}: {e}")
                    continue
            
            # معالجة الأسطر
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if file_type == 'list':
                    word = diacritics_handler.remove_diacritics(line)
                    # تطبيع النص القرآني
                    word = diacritics_handler.normalize_quranic_text(word)
                    all_words.add(word)
                else:  # text
                    for m in arabic_word_re.findall(line):
                        word = diacritics_handler.remove_diacritics(m)
                        if word:
                            # تطبيع النص القرآني
                            word = diacritics_handler.normalize_quranic_text(word)
                            all_words.add(word)
        except Exception as e:
            logging.warning(f"تعذّر قراءة الملف للتجميع: {fp} - {e}")
    return all_words

##################################
# دالة لتحميل الوسم
##################################
def load_tags(tags_file_path):
   tags_map = {}
   if os.path.exists(tags_file_path):
       with open(tags_file_path, 'r', encoding='utf-8') as f:
           pattern = re.compile(r'"([^"]+)"\s*=\s*"([^"]+)"')
           for line in f:
               line = line.strip()
               match = pattern.match(line)
               if match:
                   word = match.group(1)
                   tag = match.group(2)
                   tags_map[word] = tag
   return tags_map

##################################
# الدالة الرئيسية المحسّنة
##################################

def main():
   start_time = datetime.now()
   
####################################################################################################################################################################################################################################
   optional_tashkeel = False #True or False
   #optional_tashkeel = True #True or False
   match_whole_word = True #True or False
   #match_whole_word = False #True or False
   corpus_type = 'list' # list البحث في قائمة | text البحث في نص
   # تم إزالة خيار part_of_speech - البرنامج يقرأ الأسماء والأفعال معاً
   use_cache = True  # استخدام نظام الكاش
   use_database = True  # استخدام قاعدة البيانات
   generate_report = True  # توليد التقارير (سيكون نصيًا فقط)
   use_cross_validation = True  # استخدام التحقق التبادلي
   
   database_folder = r"قواعد البيانات"
   symbols_file_path = os.path.join(database_folder, "الخريطة.txt")
   corpus_folder = os.path.join(database_folder, "المدونة")
   tags_file_path = os.path.join(database_folder, "0.3 الوسم.txt")

   # قراءة أوزان الأسماء والأفعال معاً
   names_results_dir = os.path.join(database_folder, "النتائج_الأسماء")
   verbs_results_dir = os.path.join(database_folder, "النتائج_الأفعال")
   
   names_weights_file = os.path.join(database_folder, "0.3 أوزان_الأسماء.txt")
   verbs_weights_file = os.path.join(database_folder, "0.3 أوزان_الأفعال.txt")
   
   names_affixes_file = os.path.join(database_folder, "0.3 سوابق ولواحق_أسماء.txt")
   verbs_affixes_file = os.path.join(database_folder, "0.3 سوابق ولواحق_أفعال.txt")

   # تحميل الرموز
   with open(symbols_file_path, "r", encoding="utf-8") as sf:
       symbols_map = ast.literal_eval(sf.read())

   # تحميل السوابق واللواحق للأسماء
   if os.path.exists(names_affixes_file):
       with open(names_affixes_file, 'r', encoding='utf-8') as f:
           names_affixes_data = ast.literal_eval(f.read())
   else:
       logging.warning("لم يتم توفير ملف السوابق واللواحق للأسماء أو المسار غير صحيح. سيتم استخدام قوائم فارغة.")
       names_affixes_data = {'prefixes': [], 'suffixes': []}
   
   # تحميل السوابق واللواحق للأفعال
   if os.path.exists(verbs_affixes_file):
       with open(verbs_affixes_file, 'r', encoding='utf-8') as f:
           verbs_affixes_data = ast.literal_eval(f.read())
   else:
       logging.warning("لم يتم توفير ملف السوابق واللواحق للأفعال أو المسار غير صحيح. سيتم استخدام قوائم فارغة.")
       verbs_affixes_data = {'prefixes': [], 'suffixes': []}

   # تحميل الوسم
   tags_map = load_tags(tags_file_path)

   # إنشاء المكونات
   cache_manager = CacheManager() if use_cache else None
   db_manager = DatabaseManager() if use_database else None
   pattern_ranker = PatternRanker(db_manager) if use_database else None
   cross_validator = CrossValidator() if use_cross_validation else None
   batch_processor = BatchProcessor()
   memory_optimizer = MemoryOptimizer()
   
   # إنشاء FileManager مؤقت لقراءة الأوزان
   temp_file_manager = FileManager(
        corpus_type=corpus_type, 
        match_whole_word=match_whole_word, 
        affixes_data=names_affixes_data,  # سيتم تحديثه لكل نوع
        tags_map=tags_map,
        db_manager=None,
        cache_manager=None,
        pattern_ranker=None,
        cross_validator=None
    )
    
   # تحميل أوزان الأسماء
   names_weights = temp_file_manager.read_weights_and_derived_words(names_weights_file)
   
   # تحميل أوزان الأفعال
   temp_file_manager.affixes_data = verbs_affixes_data  # تحديث السوابق واللواحق للأفعال
   verbs_weights = temp_file_manager.read_weights_and_derived_words(verbs_weights_file)
   
   # دمج الأوزان
   all_weights = {**names_weights, **verbs_weights}

   file_paths = [os.path.join(corpus_folder, file_name) 
                 for file_name in os.listdir(corpus_folder) 
                 if file_name.endswith('.txt')]

   # جمع كلمات المدونة قبل المعالجة لحساب التغطية لاحقًا
   all_corpus_words = collect_corpus_words(file_paths, corpus_type=corpus_type)

    # معالجة الأوزان
   print(f"\n{'='*60}")
   print(f"بدء معالجة {len(all_weights)} وزن صرفي (أسماء وأفعال)...")
   print(f"{'='*60}\n")
   
   # إنشاء مهام للأسماء
   names_tasks = [(weight, derived_weights, file_paths, names_results_dir, corpus_type, match_whole_word, names_affixes_data, tags_map, symbols_map, optional_tashkeel, use_cross_validation) 
                  for weight, derived_weights in names_weights.items()]
   
   # إنشاء مهام للأفعال
   verbs_tasks = [(weight, derived_weights, file_paths, verbs_results_dir, corpus_type, match_whole_word, verbs_affixes_data, tags_map, symbols_map, optional_tashkeel, use_cross_validation) 
                  for weight, derived_weights in verbs_weights.items()]
   
   # دمج المهام
   all_tasks = names_tasks + verbs_tasks

   # معالجة الأوزان وجمع النتائج
   all_processing_results = []
   with concurrent.futures.ProcessPoolExecutor(max_workers=multiprocessing.cpu_count()) as executor:
       all_processing_results = list(tqdm(
           executor.map(process_weight, all_tasks),
           total=len(all_tasks),
           desc="معالجة الأوزان",
           unit="وزن"
       ))
   
   # حفظ النتائج في قاعدة البيانات
   if db_manager:
       print(f"\n{'='*60}")
       print("💾 حفظ النتائج في قاعدة البيانات...")
       print(f"{'='*60}\n")
       
       diacritics_handler = DiacriticsHandler()
       word_splitter = WordSplitter(diacritics_handler)
       
       for result_data in tqdm(all_processing_results, desc="حفظ في قاعدة البيانات"):
           weight = result_data['weight']
           results = result_data['results']
           
           # حساب عدد أحرف الزيادة للوزن
           extra_chars_count = sum(1 for c in weight if c in FileManager.EXTRA_CHARS)
           
           # تحديد نوع الوزن بناءً على وجوده في أوزان الأسماء أو الأفعال
           if weight in names_weights:
               pattern_type = 'اسم'
           elif weight in verbs_weights:
               pattern_type = 'فعل'
           else:
               pattern_type = 'غير محدد'  # في حالة وجود تداخل
           
           # إدراج الوزن في قاعدة البيانات
           db_manager.insert_pattern(weight, pattern_type, extra_chars_count)
           
           # حفظ النتائج
           for prefix, root, suffix in results:
               matched_word = prefix + root + suffix
               
               # تحليل الكلمة
               prefix_morph, intermediate_morph, root_morph, suffix_morph = word_splitter.split_word(
                   weight, root
               )
               root_without_diacritics = diacritics_handler.remove_diacritics(root_morph)
               
               # حساب النقاط
               score = 0
               if pattern_ranker:
                   score = pattern_ranker.calculate_score(
                       weight, matched_word, prefix, suffix, 1
                   )
               
               # حفظ في قاعدة البيانات
               db_manager.insert_result(
                   matched_word, root_without_diacritics, weight,
                   prefix, suffix, intermediate_morph, score
               )

   # حفظ الكاش
   if cache_manager:
       cache_manager.save_cache()
       logging.info("تم حفظ الكاش")

   # توليد التقارير (نصي + ملفات التغطية) وإلغاء HTML/Excel القديمين
   if generate_report and db_manager:
       report_generator = ReportGenerator(db_manager)
       stats = db_manager.get_statistics()

       # حساب مجموعة الكلمات المتعرّف عليها من نتائج المعالجة
       diacritics_handler = DiacriticsHandler()
       recognized_words = set()
       for result_data in all_processing_results:
           for prefix, root, suffix in result_data['results']:
               matched_word = prefix + root + suffix
               matched_word = diacritics_handler.remove_diacritics(matched_word)
               if matched_word:
                   recognized_words.add(matched_word)

       coverage_info = report_generator.generate_coverage_outputs(
           all_words_set=all_corpus_words,
           recognized_set=recognized_words
       )

       # إضافة معلومات الوقت
       processing_time = (datetime.now() - start_time).total_seconds()
       stats['processing_time'] = processing_time

       txt_report = report_generator.generate_text_report(stats, coverage=coverage_info, output_file="report.txt")

       print(f"\n{'='*60}")
       print("📊 ملخص الإحصائيات:")
       print(f"{'='*60}")
       print(f"إجمالي النتائج: {stats['total_results']:,}")
       print(f"الكلمات الفريدة: {stats['unique_words']:,}")
       print(f"عدد الأوزان: {stats['total_patterns']:,}")
       print(f"عدد الجذور: {stats['total_roots']:,}")
       print(f"وقت المعالجة: {processing_time:.2f} ثانية")
       print(f"\n✅ تم إنشاء التقارير:")
       print(f"   - نصي: {txt_report}")
       print(f"   - قائمة المتعرّف عليها: {coverage_info['recognized_file']}")
       print(f"   - قائمة غير المتعرّف عليها: {coverage_info['unrecognized_file']}")
       print(f"   - المخطط التفاعلي: {coverage_info['coverage_html']}")
   
   # طباعة تقرير التحقق التبادلي
   if cross_validator:
       print(f"\n{'='*60}")
       print("🔍 تقرير التحقق التبادلي:")
       print(f"{'='*60}")
       print(cross_validator.get_validation_report())
   
   # إغلاق قاعدة البيانات
   if db_manager:
       db_manager.close()
   
   print(f"\n{'='*60}")
   print(f"✅ انتهت المعالجة بنجاح!")
   print(f"{'='*60}\n")

if __name__ == "__main__":
   main()
